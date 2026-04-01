"""Universal document parser - handles xlsx, xls, csv, pdf, docx"""
import io
import tempfile
import os
import pandas as pd
import pdfplumber
import docx
from pathlib import Path
from typing import Dict, Any, Union


# ─────────────────────────────────────────────────────────────────────────────
# Public convenience function used by rfp.py routes
# ─────────────────────────────────────────────────────────────────────────────

def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Accept raw bytes + original filename, write to a temp file,
    parse it, and return the full_text string.

    This is the primary API consumed by api/routes/rfp.py.
    """
    suffix = Path(filename).suffix.lower() or ".bin"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    try:
        result = parse_document(tmp_path)
        return result.get("full_text", "")
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


# ─────────────────────────────────────────────────────────────────────────────
# Core parser — accepts a file path, returns structured dict
# ─────────────────────────────────────────────────────────────────────────────

def parse_document(file_path: str) -> Dict[str, Any]:
    """Parse any supported document type into a unified text/table structure."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix in (".xlsx", ".xls"):
        return _parse_excel(file_path)
    elif suffix == ".csv":
        return _parse_csv(file_path)
    elif suffix == ".pdf":
        return _parse_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


# ─────────────────────────────────────────────────────────────────────────────
# Private parsers
# ─────────────────────────────────────────────────────────────────────────────

def _parse_excel(file_path: str) -> Dict[str, Any]:
    xls = pd.ExcelFile(file_path)
    sheets = {}
    full_text_parts = []
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        df = df.fillna("")
        sheets[sheet_name] = df.to_dict(orient="records")
        text = f"=== Sheet: {sheet_name} ===\n"
        text += df.to_csv(index=False)
        full_text_parts.append(text)
    return {
        "type": "excel",
        "sheets": sheets,
        "full_text": "\n\n".join(full_text_parts),
    }


def _parse_csv(file_path: str) -> Dict[str, Any]:
    df = pd.read_csv(file_path).fillna("")
    return {
        "type": "csv",
        "sheets": {"Sheet1": df.to_dict(orient="records")},
        "full_text": df.to_csv(index=False),
    }


def _parse_pdf(file_path: str) -> Dict[str, Any]:
    pages = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            pages.append(text)
    full_text = "\n\n".join(pages)
    return {
        "type": "pdf",
        "sheets": {},
        "full_text": full_text,
    }


def _parse_docx(file_path: str) -> Dict[str, Any]:
    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            table_texts.append(row_text)
    full_text = "\n".join(paragraphs)
    if table_texts:
        full_text += "\n\n=== Tables ===\n" + "\n".join(table_texts)
    return {
        "type": "docx",
        "sheets": {},
        "full_text": full_text,
    }
